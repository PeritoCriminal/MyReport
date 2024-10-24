"""

Copyright (c) 2024 Marcos de Oliveira Capristo
Todos os direitos reservados.

MYREPORT é um projeto independente.
Oferece um ambiente para edição de laudos periciais,
voltado especialmente para Peritos Criminais Oficiais do Estado de São Paulo.
Idealizado e inicialmente desenvolvido pelo Perito Criminal Marcos de Oliveira Capristo.
Contato: marcos.moc@policiacientifica.sp.gov.br | (19) 9 8231-2774


"""

"""

VIEW DROGS_REPORT, A MELHORAR


"""


from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect
from django.http import HttpResponse
from ..models.drugs_report_model import DrugReport
from ..models.custom_user_model import UserRegistrationModel
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import os
import ast
import base64

from .viewbase import insert_image_from_base64_to_docx, adicionar_texto_formatado, adicionar_rodape, check_images


@login_required(login_url='/login/')
def drugs_report(request):
    user = request.user
    user_data = UserRegistrationModel.objects.get(username=user.username)


    if request.method == 'POST':  
        
        y1, m1, d1 = request.POST.get('data_atendimento').split('-')
        invert_date1 = f"{d1}-{m1}-{y1}"

        y2, m2, d2 = request.POST.get('data_liberacao').split('-')
        invert_date2 = f"{d2}-{m2}-{y2}"

        lacres_saidas = ', '.join(set(request.POST.getlist('lacre_saida[]')))
        # lacres_saidas_str = ', '.join(lacres_saidas)

        try:            
            new_report = DrugReport(
                report_number='000/00',
                city=user_data.city,
                protocol_number=request.POST.get('protocolo'),
                designated_date=invert_date1,
                exam_objective='Constatação Provisória',
                occurrence_nature='Entorpecente',
                occurring_number=request.POST.get('boletim'),
                police_station=request.POST.get('delegacia'),
                requesting_authority=request.POST.get('autoridade_requisitante'),
                activation_date=invert_date1,
                activation_time=request.POST.get('hora_atendimento'),
                service_date=invert_date2,
                service_time=request.POST.get('hora_liberacao'),
                director=user_data.director,
                nucleus=user_data.unit,
                team=user_data.team,
                reporting_expert=user.full_name,
                assistant_expert='Não se aplica',
                photographer='Não se aplica',
                draftsman='Não se aplica',
                allSubstances=request.POST.getlist('substance[]'),
                materialReceivedObservations=request.POST.get('allMaterialObservations'),
                materialImage=request.POST.get('imageGeneralBase64'),
                materialImageCaption=request.POST.get('labelOfImgOfAllMaterialReceived'),
                listOfEnvolvedPeople=request.POST.getlist('people_involved[]'),
                listOfPackagings=request.POST.getlist('packaging[]'),
                listOfMorphology=request.POST.getlist('morphology[]'),
                listOfGrossMass=[float(x) for x in request.POST.getlist('massa_bruta[]') if x],
                listOfLiquidMass=[float(x) for x in request.POST.getlist('massa_liquida[]') if x],
                listOfReturned=request.POST.get('listOfReturndCheckBoxes'),
                listOfCounterProof=request.POST.getlist('returnedOrCounterproof[]'),
                listOfEntranceSeal=request.POST.getlist('lacre_entrada[]'),
                listOfExitseal=request.POST.getlist('lacre_saida[]'),
                listOfpackagingAndMorphology=request.POST.getlist('packagingAndMorphologiObservations[]'),
                listOfResultOfExams=request.POST.getlist('resultOfExams[]'),
                examImages=request.POST.getlist('imageItemExaminatedBase64[]'),
                examImageCaptions=request.POST.getlist('labelOfImgItemExaminatedReceived[]'),                
                returnedItemsImage=request.POST.get('material-devolvidoBase64'),
                returnedItemsCaption=request.POST.get('material_devolvido_legenda'),
                counterProofImage=request.POST.get('contrapericia-imagemoBase64'),
                counterProofCaption=request.POST.get('contrapericia_legenda'),
                considerations = request.POST.get('considerations'),
                conclusion = request.POST.get('conclusao'),
            )

            """
            
            ALINHA ABAIXO ESTÁ COMENTADA PARA NÃO SALVAR O OBJETO NO BANCO DE DADOS.
            A TABELA EXISTE E O MÓDULO ESTÁ PRONTO PARA ESSA FUNCIONALIDADE, 
            MAS NÃO FOI IMPLEMENTADO A RECUPERAÇÃO DOS DADOS PARA SUA ALTERAÇÃO OU EXCLUSÃO.
                    
            """

            # new_report.save() 


            """     
            
            GERAR ARQUIVO DOCX E ENVIAR AO USUÁRIO
            
            """
            template_path = os.path.join('myreportapp', 'static', 'doctemplates', 'report.docx')

            def adicionar_itens(doc, n):

                list_of_returned = ast.literal_eval(new_report.listOfReturned.replace('true', 'True').replace('false', 'False'))

                for i in range(n):
                    if n == 1:
                        item_q = f'Item único'
                    else:
                        item_q = f'Item {i + 1}'                 
                    adicionar_texto_formatado(doc, item_q, f' (Acondicionado sob o lacre {new_report.listOfEntranceSeal[i]})')
                    paragrafo = doc.paragraphs[-1]
                    paragrafo.paragraph_format.space_before = Pt(24)
                    run_negrito = paragrafo.runs[0]
                    run_negrito.underline = True                   
                    
                    adicionar_texto_formatado(doc, 'Descrição: ', new_report.listOfpackagingAndMorphology[i], 2)
                    adicionar_texto_formatado(doc, 'Massa Bruta: ', f'{new_report.listOfGrossMass[i]} gramas', 2)
                    adicionar_texto_formatado(doc, 'Massa Líquida: ', f'{new_report.listOfLiquidMass[i]} gramas', 2)
                    adicionar_texto_formatado(doc, 'Quantidade retirada para análise e/ou contraperícia: ', new_report.listOfCounterProof[i], 2)
                    adicionar_texto_formatado(doc, 'Resultado: ', new_report.listOfResultOfExams[i], 2)

                    if list_of_returned[i]:
                        returned_if = f'O restante do item (material , invólucro(s) e lacre(s)) foi devolvido à autoridade policial requisitante nos termos das exigências legais, sob o lacre número {new_report.listOfExitseal[i]}.'
                    else:
                        returned_if = 'As embalagens e os respectivos lacres foram aqui inutilizados conforme as Portarias SPTC No 112 de 29-6-2016 e SPTC no 63 de 30 de abril de 2015.'          
                    
                    adicionar_texto_formatado(doc, '', returned_if, 2)
            
            doc = None
            try:
                doc = Document(template_path)
            except Exception as e:
                print(f'Erro ao abrir o documento: {e}')
                doc = Document()

            if doc.paragraphs and not doc.paragraphs[0].text.strip():
                p = doc.paragraphs[0]._element
                p.getparent().remove(p)

            doc.add_heading('LAUDO DE CONSTATAÇÃO', 0)
            adicionar_texto_formatado(doc, 'Registro de Entrada RE: ', new_report.protocol_number)
            adicionar_texto_formatado(doc, 'Origem: BO - ', new_report.occurring_number)
            adicionar_texto_formatado(doc, 'Autoridade Requisitante: ', new_report.requesting_authority)            

            if len(new_report.listOfEnvolvedPeople) > 1:
                formatted_people = ', '.join(new_report.listOfEnvolvedPeople)
            else:
                formatted_people = new_report.listOfEnvolvedPeople[0]
            
            adicionar_texto_formatado(doc, 'Nome(s) do(s) Envolvido(s): ', f'{formatted_people}.')
            adicionar_texto_formatado(doc, 'Data do Exame: ', new_report.designated_date)
    
            preamble_text = new_report.generate_preamble()
    
            preamble_paragraph = doc.add_paragraph(preamble_text)
            
            for run in preamble_paragraph.runs:
                run.font.size = Pt(11)
                run.font.name = 'Arial'

            number_of_itens = len(new_report.listOfPackagings)

            if number_of_itens == 1:
                sing_or_pl = 'Item'
            else:
                sing_or_pl = 'Itens'

            adicionar_texto_formatado(doc, 'Dos Materiais Recebidos e Examinados ', f'({number_of_itens} {sing_or_pl}):')

            doc.add_paragraph(new_report.materialReceivedObservations)

            doc.add_paragraph('O Exame Revelou:')

            adicionar_itens(doc, number_of_itens)
           
            adicionar_texto_formatado(doc, '', new_report.considerations)

            adicionar_texto_formatado(doc, '', 'Este laudo segue assinado digitalmente e estará arquivado no sistema GDL da Superintendência da Polícia Técnico Científica do Estado de São Paulo.')

            assinado_paragraph = doc.add_paragraph('ASSINADO DIGITALMENTE')
            assinado_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            expert_paragraph = doc.add_paragraph(new_report.reporting_expert)
            expert_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            perito_paragraph = doc.add_paragraph('Perito(a) Criminal')
            perito_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            hasImg = check_images(new_report.materialImage, new_report.examImages, new_report.returnedItemsImage, new_report.counterProofImage)
            
            adicionar_texto_formatado(doc, hasImg, '')

            num_img_caption = 0
            
            num_img_caption = insert_image_from_base64_to_docx(doc, new_report.materialImage, new_report.materialImageCaption, num_img_caption)

            for i in range(len(new_report.examImages)):
                num_img_caption = insert_image_from_base64_to_docx(doc, new_report.examImages[i], new_report.examImageCaptions[i], num_img_caption)
    
            num_img_caption = insert_image_from_base64_to_docx(doc, new_report.returnedItemsImage, new_report.returnedItemsCaption, num_img_caption)

            num_img_caption = insert_image_from_base64_to_docx(doc, new_report.counterProofImage, new_report.counterProofCaption, num_img_caption)

            adicionar_rodape(doc, f'RE: {new_report.protocol_number} | Boletim {new_report.occurring_number} - {new_report.police_station}')

            if lacres_saidas:
                new_section = doc.add_section(WD_SECTION.NEW_PAGE)

                new_section.header.is_linked_to_previous = True

                new_section.footer.is_linked_to_previous = False
                for paragraph in new_section.footer.paragraphs:
                    paragraph.clear()
                doc.add_paragraph('')
                doc.add_heading('DEVOLUÇÃO DE MATERIAL EXAMINADO', 0)
                if ',' in lacres_saidas:
                    paragraph = doc.add_paragraph('Lacres: ')
                else:
                    paragraph = doc.add_paragraph('Lacre:')    
                run = paragraph.add_run(lacres_saidas)
                doc.add_paragraph('')
                doc.add_paragraph(f'Origem: Boletim {new_report.occurring_number} | {new_report.police_station}')
                doc.add_paragraph(f'Autoridade Requisitante: {new_report.requesting_authority}')
                doc.add_paragraph(f'Registro de Entrada RE: {new_report.protocol_number}')
                doc.add_paragraph(f'Perito: {new_report.reporting_expert}')
                run.bold = True
                doc.add_paragraph('')
                adicionar_texto_formatado(doc, 'Responsavel Pela Retirada','')
                doc.add_paragraph('Nome: ___________________________________________________')
                doc.add_paragraph('RG:  ________________________________')
                doc.add_paragraph('Data da Retirada: _______/_______/_______')
                doc.add_paragraph('')
                adicionar_texto_formatado(doc, 'Assinatura: ', '_________________________________________')            
            
            
            doc_buffer = BytesIO()  
            doc.save(doc_buffer)
            doc_buffer.seek(0)

            file_name = f"{new_report.protocol_number.replace('/', '_')}$narc.docx"

            response = HttpResponse(doc_buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename={file_name}'
            
            return response
        except Exception as e:
            return HttpResponse(f'Erro ao criar relatório: {str(e)}')
    now = datetime.now()
    minutes = (now.minute // 15) * 15
    rounded_time = now.replace(minute=minutes, second=0, microsecond=0)
    formatted_current_time = rounded_time.strftime('%H:%M')
    before_time = rounded_time - timedelta(hours=1.5)
    formatted_before_time = before_time.strftime('%H:%M')
    today_date = now.date()
    if before_time.day != now.day:
        before_date = today_date - timedelta(days=1)  
    else:
        before_date = today_date  
    msg_about_this_form_to_user = (
        f'{user_data.full_name},<br>Este formulário serve como uma alternativa '
        'caso o Sistema GDL não esteja disponível. Preencha todos os campos com atenção, '
        'verifique todos os dados cuidadosamente, salve o laudo em formato PDF e '
        'assine-o com seu certificado digital. Caso o laudo seja enviado à autoridade requisitante, '
        'ele deverá ser registrado oportunamente no Sistema GDL.'
    )

    packaging = {
        'Selecione uma opção': '',
        'Plástico com nó': 'invólucro(s) plástico(s) fechado(s) por nó encerrando',
        'Plástico com filme': 'invólucro(s) plástico(s) do tipo "filme" retorcido encerrando',
        'Plástico com zip': 'invólucro(s) plástico(s) fechado(s) por pressão (tipo "zip") encerrando',
        'Plástico com calor': 'invólucro(s) plástico(s) fechado(s) por aquecimento encerrando',
        'Plástico com alumínio': 'invólucro(s) constituído(s) por plástico e papel alumínio encerrando',
        'Fita adesiva': 'invólucro(s) constituído(s) por fita(s) adesiva(s) retorcida(s) encerrando',
        'Eppendorf': 'microtubo(s) plástico(s) do tipo "Eppendorf" dotado(s) de tampa própria encerrando',
        'Frasco vítreo': 'Frasco(s) vítreo(s) fechado(s) por batoque e tampa própria contendo',
        'Invólucro geral': 'invólucro(s) plástico(s) encerrando',
        'Outros': '[DESCREVA A EMBALAGEM]',
        'Apenas papel alumínio': 'invólucro(s) constituído(s) por papel alumínio encerrando',
        'Frasco com válvula (lança perfume)': 'Frasco(s) vítreos dotado(s) de válvula aspersora contendo',
        'Frasco plástico': 'Frasco(s) plástico(s) fechado(s) por tampa própria contendo',
        'Fita adesiva com plástico filme': 'invólucro(s) constituído(s) por fita(s) adesiva(s) e plástico(s) encerrando',
        'Garrafa pet': 'Garrafa(s) aparentemente do tipo PET, fechada(s) por tampa própria de rosca, contendo'
    }

    morphology = {
        'Selecione uma opção': '',
        'erva': 'porção de fragmentos vegetais ressequidos, constituídos de folhas, folíolos, inflorescências, caules e frutos.',
        'tijolo': 'contendo porção de fragmentos vegetais ressequidos, constituídos de folhas, folíolos, inflorescências, caules e frutos, compactados na forma de tijolo.',
        'planta': 'contendo planta arbustiva com [XX] cm de comprimento e constituído por raiz, caule, folhas, folíolos, inflorescência e frutos',
        'haxixe': 'contendo porção de substância de aspecto resinoso de coloração amarronzada e de morfologia irregular',
        'cigarro íntegro': 'contendo cigarro(s) artesanal(ais) confeccionado(s) em papel, contendo fragmentos vegetais ressequidos, constituídos de folhas, folíolos, inflorescências, caules e frutos.',
        'cigarro queimado': 'contendo cigarro(s) artesanal(ais) parcialmente queimado(s), confeccionado(s) em papel, contendo fragmentos vegetais ressequidos, constituídos de folhas, folíolos, inflorescências, caules e frutos.',
        'pó': 'contendo porção de material sólido particulado',
        'pedra': 'contendo porção de material sólido petrificado',
        'líquido': 'contendo fração líquida translúcida e volátil.',
        'resina': 'contendo porção de material resinoso',
        'comprimido': 'contendo porção de material particulado, compactado na forma de comprimido',
        'selo': 'contendo segmento de papel ilustrado, do tipo "picote" (selo).',
        'outros': '[Descreva a Morfologia]',
        'granulado': 'contendo porção de material sólido particulado com grânulos',
        'dichavador': 'contendo dichavador, do tipo triturador de rotação manual, da marca [X]/ sem marca aparente, de [material: madeira/metal/plástico/material sintético], composto por [X] partes, apresentado resquícios/sujidades.',
        'Semente/ fruto de maconha': 'contendo porção de material com aspecto de origem vegetal, de formato esférico, coloração amarronzada e com diâmetro médio de aproximadamente 3,0 mm (três milímetros).',
        'K4': 'contendo segmento(s) de papel de cor [descreva a cor]',
        'Liquidificador': 'contendo liquidificador, da marca [X]/ sem marca aparente, composto por base e recipiente com hélice de pás cortantes, apresentando resquícios/sujidades.',
        'balança': 'contendo balança, da marca [X]/ sem marca aparente, com mostrador digital, de carga máxima nominal de [X], apresentando resquícios/sujidades.',
        'faca': 'contendo faca, da marca [X]/ sem marca aparente, constituída por uma lâmina metálica cortante presa a um cabo de [material: madeira/metal/plástico], apresentando resquícios/sujidades.',
        'tesoura': 'contendo tesoura, da marca [X]/ sem marca aparente, constituída por duas lâminas metálicas cortantes, unidas por um eixo, apresentando resquícios/sujidades.',
        'embalagem plástica com resquícios': 'contendo embalagens plásticas com massa total de [X] gramas e resquícios do material [descrição do material].',
        'embalagem plástica': 'contendo embalagem plástica com massa total de [X] gramas e [descrição do material].'
    }
    exam_resulting = {
        'Selecione uma opção': '',
        'Negativo': 'NÃO FOI POSSÍVEL IDENTIFICAR presença de substâncias elencadas nas listas A, B e F da Portaria SVS/MS 344/98 e atualizações posteriores, ou na Portaria MJSP 204/2022, em sua lista III, conforme a(s) técnica(s) utilizada(s) (Portaria SPTC 42/2024).',
        'Inconclusivo': 'Os exames/análises preliminares mostraram-se INCONCLUSIVOS, sendo necessárias análises mais complexas e morosas, incompatíveis com a rapidez demandada pelos exames de constatação. O resultado deste presente item seguirá em laudo definitivo.',
        'Cocaina': 'foi DETECTADA presença da substância COCAÍNA, constante na lista F1 da Portaria SVS/MS 344/98 e atualizações posteriores',
        'Crack': 'foi DETECTADA presença da substância COCAÍNA, constante na lista F1 da Portaria SVS/MS 344/98 e atualizações posteriores',
        'Maconha': 'foi DETECTADA presença da substância TETRAHIDROCANNABINOL (THC), constante na lista F2 da Portaria SVS/MS 344/98 e atualizações posteriores',
        }
    
    conterproof = {
        'ff':'todo o material disponível foi utilizado nas análises do presente Laudo. Face a sua exiguidade não foi possível armazenar amostra de contraperícia.',
        'fv':'todo o material foi aqui retirado para análises, sendo o remanescente destas análises armazenado sob a forma de contraperícia.',
        'vv':'uma amostra de aproximadamente 2 g (dois gramas) foi aqui retirada para análises, sendo o remanescente destas análises armazenado sob a forma de contraperícia.'
    }

    context = {
        'protocol_prefix': 'TOX',
        'delegacia': 'Del. Sec. Limeira Plantão',
        'lacre_saida': '',
        'today_date': today_date,
        'current_time': formatted_current_time,
        'material': 'Todo material recebido encontrava-se acondicionado em invólucro(s) plástico(s) lacrado(s), acompanhado da requisição de exame pericial.',
        'before_time': formatted_before_time,
        'before_date': before_date,
        'user_name': user.full_name,
        'msg_about_this_form_to_user': msg_about_this_form_to_user,
        'packaging': packaging,
        'morphology': morphology,
        'exam_resulting': exam_resulting,
        'conterproof': conterproof,
        'drUser': user_data.full_name,
        'alert': '(Este Laudo é de caráter provisório e não confirma necessariamente o resultado da identificação que será enviado no Laudo Definitivo)',
    } 

    return render(request, 'drugs_report.html', context)
