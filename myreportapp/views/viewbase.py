"""

Copyright (c) 2024 Marcos de Oliveira Capristo
Todos os direitos reservados.

MYREPORT é um projeto independente.
Oferece um ambiente para edição de laudos periciais,
voltado especialmente para Peritos Criminais Oficiais do Estado de São Paulo.
Idealizado e inicialmente desenvolvido pelo Perito Criminal Marcos de Oliveira Capristo.
Contato: marcos.moc@policiacientifica.sp.gov.br | (19) 9 8231-2774


"""

"""

FUNÇÕES py COMUNS A TODOS OS MÓDULOS, EM DESENVOLVIMENTO


"""

import base64
from io import BytesIO
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import ast
import re 

# Variável global para contar as imagens
img_label = 0

def insert_image_from_base64_to_docx(doc, base64_string, label, num_img, width_cm=12):
    global img_label  # Declarar a variável como global

    if not base64_string:
        return num_img

    try:
        # Remove o prefixo 'data:image/png;base64,' ou similar, se existir
        if ';base64,' in base64_string:
            base64_string = base64_string.split(';base64,')[1]

        # Decodifica a string Base64 em bytes
        image_data = base64.b64decode(base64_string)

        # Salva a imagem temporariamente em um objeto BytesIO
        image_stream = BytesIO(image_data)

        # Insere um parágrafo e adiciona a imagem a ele
        paragraph = doc.add_paragraph()  # Cria um novo parágrafo
        run = paragraph.add_run()         # Adiciona um run ao parágrafo
        run.add_picture(image_stream, width=Cm(width_cm))  # Adiciona a imagem
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centraliza o parágrafo que contém a imagem

        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(24)
        paragraph_format.space_after = Pt(0)
        paragraph_format.keep_with_next = True

        num_img += 1;
        img_label += 1  # Incrementa a contagem de imagens
        legenda = f'Imagem {num_img} - {label}'

        caption_paragraph = doc.add_paragraph()
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        caption_format = caption_paragraph.paragraph_format
        caption_format.space_before = Pt(0)
        
        # Adiciona a legenda como um run com formatação
        caption_run = caption_paragraph.add_run(legenda)
        caption_run.italic = True  # Define o texto como itálico
        caption_run.font.size = Pt(11) 

        return num_img

    except Exception as e:
        return f"An error occurred: {e}"
    




def adicionar_texto_formatado(doc, texto_negrito, texto_normal, recuo=1):
            """
            Adiciona um texto formatado ao documento.

            :param doc: O documento ao qual o texto será adicionado.
            :param texto_negrito: O texto que deve ser adicionado em negrito.
            :param texto_normal: O texto que deve ser adicionado normalmente.
            :param recuo: O recuo à esquerda do parágrafo em centímetros. Padrão é 1 cm.
            """
            # Cria um parágrafo e adiciona o texto em negrito
            paragrafo = doc.add_paragraph()
            
            # Define o recuo à esquerda
            paragrafo.paragraph_format.left_indent = Cm(recuo)

            # Adiciona o texto em negrito
            run_negrito = paragrafo.add_run(texto_negrito)
            run_negrito.bold = True
            
            # Adiciona o texto normal em seguida
            paragrafo.add_run(texto_normal)

        # print(f'Path do template: {os.path.abspath(template_path)}')






def adicionar_rodape(doc, texto_rodape):
        # Adicionar o rodapé ao documento
        section = doc.sections[-1]
        footer = section.footer

        # Criar um parágrafo para o rodapé
        paragrafo_rodape1 = footer.paragraphs[0]
        paragrafo_rodape1.text = texto_rodape
        paragrafo_rodape1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Definir a fonte e o tamanho do rodapé
        run_rodape1 = paragrafo_rodape1.runs[0]
        run_rodape1.font.name = 'Times New Roman'
        run_rodape1.font.size = Pt(10)

        # Adicionar a segunda linha: Página X de Y
        paragrafo_rodape2 = footer.add_paragraph()
        paragrafo_rodape2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Texto padrão antes da numeração da página
        run_rodape2 = paragrafo_rodape2.add_run("Página ")
        run_rodape2.font.name = 'Times New Roman'
        run_rodape2.font.size = Pt(10)

        # Definir espaçamento de 6 pontos antes e depois, e espaçamento de linha simples
        paragrafo_rodape1.paragraph_format.space_before = Pt(12)
        paragrafo_rodape1.paragraph_format.space_after = Pt(0)
        paragrafo_rodape1.paragraph_format.line_spacing = 1  # Espaçamento simples

        # Adicionando a numeração de página
        field_code = 'PAGE'
        page_field = OxmlElement('w:fldSimple')
        page_field.set(qn('w:instr'), field_code)
        run_num_page = OxmlElement('w:r')
        run_text = OxmlElement('w:t')
        run_text.text = "1"  # Placeholder que será substituído por Word
        run_num_page.append(run_text)
        page_field.append(run_num_page)
        paragrafo_rodape2._element.append(page_field)

        # Adicionar o texto " de "
        run_de = paragrafo_rodape2.add_run(" de ")
        run_de.font.name = 'Times New Roman'
        run_de.font.size = Pt(10)

        paragrafo_rodape2.paragraph_format.space_before = Pt(0)
        paragrafo_rodape2.paragraph_format.space_after = Pt(6)
        paragrafo_rodape2.paragraph_format.line_spacing = 1  # Espaçamento simples

        # Adicionando a contagem total de páginas
        total_page_field_code = 'NUMPAGES'
        total_page_field = OxmlElement('w:fldSimple')
        total_page_field.set(qn('w:instr'), total_page_field_code)
        run_total_pages = OxmlElement('w:r')
        run_total_text = OxmlElement('w:t')
        run_total_text.text = "1"  # Placeholder que será substituído por Word
        run_total_pages.append(run_total_text)
        total_page_field.append(run_total_pages)
        paragrafo_rodape2._element.append(total_page_field)





def check_images(materialImage, examImages, returnedItemsImage, counterProofImage):
        """
        Verifica se há imagens entre os atributos fornecidos.

        :param materialImage: A string Base64 da imagem principal.
        :param examImages: A string que representa a lista de strings Base64 de imagens examinadas.
        :param returnedItemsImage: A string Base64 da imagem de itens devolvidos.
        :param counterProofImage: A string Base64 da imagem de contraprova.
        :return: Uma string informando se há ou não imagens.
        """
        # Verifica se a imagem principal não é vazia
        hasImg = 'Ilustrando as peças do exame:'

        if materialImage:
            return hasImg

        # Tenta avaliar a lista de imagens examinadas
        try:
            exam_images_list = ast.literal_eval(examImages)  # Converte a string para lista
        except (ValueError, SyntaxError):
            exam_images_list = []  # Se não for possível converter, assume lista vazia

        # Verifica se há imagens na lista de imagens examinadas
        if any(image.strip() for image in exam_images_list):  # Verifica se há algum item não vazio
            return hasImg

        # Verifica se a imagem de itens devolvidos não é vazia
        if returnedItemsImage:
            return hasImg

        # Verifica se a imagem de contraprova não é vazia
        if counterProofImage:
            return hasImg

        return "Não há imagens disponíveis."



def format_filename(string):
    string = string.strip()
    string = re.sub(r'[^a-z ]', '', string)
    string = string.replace(' ', '_')    
    return string