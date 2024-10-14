"""

Copyright (c) 2024 Marcos de Oliveira Capristo
Todos os direitos reservados.

MYREPORT é um projeto independente.
Oferece um ambiente para edição de laudos periciais,
voltado especialmente para Peritos Criminais Oficiais do Estado de São Paulo.
Idealizado e inicialmente desenvolvido pelo Perito Criminal Marcos de Oliveira Capristo.
Contato: marcos.moc@policiacientifica.sp.gov.br | (19) 9 8231-2774


"""

"""

VIEW FURTO, A DESENVOLVER


"""


from django.shortcuts import render, redirect, get_object_or_404, HttpResponse
from ..models.custom_user_model import UserRegistrationModel
from ..models.theft_report_model import TheftReportModel
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

from .viewbase import adicionar_rodape, insert_image_from_base64_to_docx, format_filename, adicionar_texto_formatado, format_date, format_hour

def theft_report_view(request, report_id=None):
    user = request.user
    user_data = UserRegistrationModel.objects.get(username=user.username)
    message = None
    numberOfLocals = 0
    local_preservation = ''
    list_local_subtitle = []
    list_local_description = []
    list_local_img = []
    list_local_label = []
    #new_report = False   # Define como falso inicialmente

    # Verificar se o relatório já existe (edição) ou se será um novo
    if report_id:
        report = get_object_or_404(TheftReportModel, id=report_id)
        numberOfLocals = len(report.localsubtitle)  
        list_local_subtitle = report.localsubtitle 
        local_preservation = report.preservation_context
        list_local_description = report.localdescription
        list_local_img = report.localimgbase64
        list_local_label = report.locallegend
        message = f'Atualização do Laudo {report.report_number}'  
    else:
        report = TheftReportModel()  
        message = 'Elaboração de novo laudo.'
        

    if request.method == 'POST':
        rep_id = request.POST.get('report_id')
        try:
            # Se for uma edição de relatório existente
            if rep_id:  # Edita o relatório existente
                report = get_object_or_404(TheftReportModel, id=rep_id)
                report.report_number = request.POST.get('laudo')
                report.protocol_number = request.POST.get('protocolo')
                report.designated_date = request.POST.get('data_designacao')
                report.exam_objective = request.POST.get('objetivo')
                report.occurrence_nature = request.POST.get('naturesa')
                report.occurring_number = request.POST.get('boletim')
                report.police_station = request.POST.get('delegacia')
                report.requesting_authority = request.POST.get('autoridade_requisitante')
                report.activation_date = request.POST.get('data_acionamento')
                report.activation_time = request.POST.get('hora_acionamento')
                report.service_date = request.POST.get('data_atendimento')
                report.service_time = request.POST.get('hora_atendimento')
                report.photographer = request.POST.get('fotografo')
                report.preservation_context = request.POST.get('preservationContext')
                report.localsubtitle = request.POST.getlist('local-subtitle[]')
                report.localdescription = request.POST.getlist('local-description[]')
                report.localimgbase64 = request.POST.getlist('local-img-to-text-base64[]')
                report.locallegend = request.POST.getlist('label-local-img[]')
                report.considerations = request.POST.get('considerations')
                report.conclusion = request.POST.get('conclusions')
                report.save()  # Salva o relatório editado
                message += " - Relatório atualizado com sucesso!"
            
            # Se não for uma edição, cria um novo relatório
            else:
                newReport = TheftReportModel(
                    report_number=request.POST.get('laudo'),
                    city=user_data.city,
                    protocol_number=request.POST.get('protocolo'),
                    designated_date=request.POST.get('data_designacao'),
                    exam_objective=request.POST.get('objetivo'),
                    occurrence_nature=request.POST.get('naturesa'),
                    occurring_number=request.POST.get('boletim'),
                    police_station=request.POST.get('delegacia'),
                    requesting_authority=request.POST.get('autoridade_requisitante'),
                    activation_date=request.POST.get('data_acionamento'),
                    activation_time=request.POST.get('hora_acionamento'),
                    service_date=request.POST.get('data_atendimento'),
                    service_time=request.POST.get('hora_atendimento'),
                    director=user_data.director,
                    nucleus=user_data.unit,
                    team=user_data.team,
                    reporting_expert=user.full_name,
                    assistant_expert='Não se aplica',
                    photographer=request.POST.get('fotografo'),
                    draftsman='Não se aplica',
                    preservation_context=request.POST.get('preservationContext'),
                    localsubtitle=request.POST.getlist('local-subtitle[]'),
                    localdescription=request.POST.getlist('local-description[]'),
                    localimgbase64=request.POST.getlist('local-img-to-text-base64[]'),
                    locallegend=request.POST.getlist('label-local-img[]'),
                    considerations=request.POST.get('considerations'),
                    conclusion=request.POST.get('conclusions'),
                )
                newReport.save()  # Salva o novo relatório
                message += " - Novo relatório salvo com sucesso!"
                
            return redirect('user_reports')

        except Exception as e:
            print(f'Erro: {e}')
            message = f"Erro ao salvar o relatório: {e}"

    context = {
        'local_preservation': local_preservation,
        'list_local_subtitle': list_local_subtitle,
        'list_local_description': list_local_description,
        'list_local_img': list_local_img,
        'list_local_label': list_local_label,
        'number_of_locals': numberOfLocals,
        'user_data': user_data,
        'report': report,
        'msg_about_this_form_to_user': 'Editor Laudo',
        'message': message,
    }

    return render(request, 'furto.html', context)


def generate_theft_docx(request, report_id):
    # Obtém o relatório com base no ID
    theft_report = get_object_or_404(TheftReportModel, pk=report_id)

    # Cria o documento DOCX
    template_path = os.path.join('myreportapp', 'static', 'doctemplates', 'report.docx')

    doc = None
    try:
        doc = Document(template_path)
    except Exception as e:
        print(f'Erro ao abrir o documento: {e}')
        doc = Document()

    if doc.paragraphs and not doc.paragraphs[0].text.strip():
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)
    
    # Adiciona o título
    doc.add_heading(f'Laudo {theft_report.report_number}', 0)
    
    # Adiciona os campos do relatório
    
    preamble_text = theft_report.generate_preamble()
    
    preamble_paragraph = doc.add_paragraph(preamble_text)
    
    for run in preamble_paragraph.runs:
        run.font.size = Pt(11)
        run.font.name = 'Arial'

    doc.add_heading('Dados da Requisição de Exame', 1)
    doc.add_paragraph(f'Autoridade Requisitante: {theft_report.requesting_authority}')
    doc.add_paragraph(f'Boletim: {theft_report.occurring_number} - {theft_report.police_station}')
    doc.add_paragraph(f'Objetivo: {theft_report.exam_objective}')
    doc.add_paragraph(f'Natureza da Ocorrência: {theft_report.occurrence_nature}')
    doc.add_heading('Histórico do Atendimento', 1)
    # doc.add_paragraph(f'Número do Laudo: {theft_report.report_number}')
    doc.add_paragraph(f'Registro de Entrada: {theft_report.protocol_number}')
    doc.add_paragraph(f'Data e hora do acionamento: {format_date(theft_report.activation_date)} às {format_hour(theft_report.activation_time)}.')
    doc.add_paragraph(f'Data e hora do atendimento: {format_date(theft_report.service_date)} às {format_hour(theft_report.service_time)}.')
    doc.add_paragraph(f'Perito: {theft_report.reporting_expert}')
    doc.add_paragraph(f'Fotografia e apoio técnico: {theft_report.photographer}')
    doc.add_heading(f'Descrição e Exame do Local', 1)

    doc.add_heading(f'Preservação', 2)
    preservation = theft_report.preservation_context
    if preservation:
        preservation = preservation.replace('\r', '').strip()  # Remove caracteres de controle
        paragraphs = preservation.split('\n')
        for paragraph in paragraphs:
            doc.add_paragraph(paragraph)
    
    counter = len(theft_report.localsubtitle)
    num_img = 0

    for i in range(counter):
        subtitle = theft_report.localsubtitle[i]
        if subtitle:
            doc.add_heading(subtitle, 2)
        description = theft_report.localdescription[i]
        if description:
            description = description.replace('\r', '').strip()  # Remove caracteres de controle
            paragraphs = description.split('\n')
            for paragraph in paragraphs:
                doc.add_paragraph(paragraph)
        img = theft_report.localimgbase64[i]
        label = theft_report.locallegend[i]
        if img:
            num_img = insert_image_from_base64_to_docx(doc, img, label, num_img)

    considerations = theft_report.considerations
    if considerations:
        doc.add_heading('Considerações', 1)
        considerations = considerations.replace('\r', '').strip()
        paragraphs = considerations.split('\n')
        for paragraph in paragraphs:
            doc.add_paragraph(paragraph)

    conclusion = theft_report.conclusion
    if conclusion:
        doc.add_heading('Considerações', 1)
        conclusion = conclusion.replace('\r', '').strip()
        paragraphs = conclusion.split('\n')
        for paragraph in paragraphs:
            doc.add_paragraph(paragraph)

    doc.add_paragraph('')
    adicionar_texto_formatado(doc, '', 'Este laudo segue assinado digitalmente e encontra-se arquivado no sistema GDL da Superintendência da Polícia Técnico Científica do Estado de São Paulo.', 0)

    # Criar o parágrafo assinado digitalmente
    assinado_paragraph = doc.add_paragraph('ASSINADO DIGITALMENTE')
    assinado_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Definir a fonte e o tamanho para 11
    run_assinado = assinado_paragraph.runs[0]
    run_assinado.font.size = Pt(11)
    run_assinado.font.name = 'Times New Roman'  # ou outra fonte de sua preferência

    # Ajustar espaçamento entre parágrafos
    assinado_paragraph.paragraph_format.space_before = Pt(24)
    assinado_paragraph.paragraph_format.space_after = Pt(6)

    # Criar o parágrafo com o nome do perito
    expert_paragraph = doc.add_paragraph(theft_report.reporting_expert)
    expert_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Definir a fonte e o tamanho para 11
    run_expert = expert_paragraph.runs[0]
    run_expert.font.size = Pt(11)
    run_expert.font.name = 'Times New Roman'

    # Ajustar espaçamento entre parágrafos
    expert_paragraph.paragraph_format.space_before = Pt(6)
    expert_paragraph.paragraph_format.space_after = Pt(6)

    # Criar o parágrafo Perito Criminal
    perito_paragraph = doc.add_paragraph('Perito(a) Criminal')
    perito_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Definir a fonte e o tamanho para 11
    run_perito = perito_paragraph.runs[0]
    run_perito.font.size = Pt(11)
    run_perito.font.name = 'Times New Roman'

    # Ajustar espaçamento entre parágrafos
    perito_paragraph.paragraph_format.space_before = Pt(6)
    perito_paragraph.paragraph_format.space_after = Pt(6)
    
    adicionar_rodape(doc, f'Laudo: {theft_report.report_number} | Boletim: {theft_report. occurring_number} - {theft_report.police_station}')

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    report_number_formatted = theft_report.report_number.replace('/', '_').replace('.', '')
    occurrence_nature_formatted = format_filename(theft_report.occurrence_nature)
    file_name = f"{report_number_formatted}${occurrence_nature_formatted}.docx"

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename={file_name}'
    
    doc.save(response)

    return response